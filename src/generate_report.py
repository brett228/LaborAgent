import datetime
from pathlib import Path

from jinja2 import Template
from docxtpl import DocxTemplate
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
import pypandoc
pypandoc.download_pandoc()


# -------------------------
# 0) URL / 환경 설정
# -------------------------
try:
    BASE_DIR = Path(__file__).resolve().parent.parent
except NameError:
    BASE_DIR = Path.cwd()

TEMPLATE_PATH = BASE_DIR / "templates"


# ---------------------------------------------------------
# 1. Markdown Template Load & Render
# ---------------------------------------------------------
def render_markdown(template_path: str, output_path: str, context: dict):
    """Render markdown from a Jinja2 template."""

    with open(template_path, "r", encoding="utf-8") as f:
        template = Template(f.read())

    rendered = template.render(**context)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(rendered)

    return output_path

# ---------------------------------------------------------
# 2. Markdown → PDF 변환
# ---------------------------------------------------------
def convert_md_to_pdf(md_path: str, pdf_path: str):
    """Convert markdown to PDF using pypandoc."""
    pypandoc.convert_file(
        md_path,
        to="pdf",
        format="md",
        outputfile=pdf_path,
        extra_args=["--standalone", "--pdf-engine=wkhtmltopdf"]
    )
    return pdf_path

# ---------------------------------------------------------
# 3. Combined Runner
# ---------------------------------------------------------
if __name__ == "__main__":
    TEMPLATE = "templates/consult_template.md"
    MD_OUTPUT = "opinion.md"
    PDF_OUTPUT = "opinion.pdf"

    context = {
        "logo_path": "images/logo.png",
        "query_summary": "사용자 질의 요약 내용...",
        "related_laws": "관련 법령 조문 정리...",
        "related_cases": "관련 판례 및 민원 사례 요약...",
        "answer": "법령 및 사례를 바탕으로 한 종합 답변..."
    }

    print("Generating Markdown...")
    render_markdown(TEMPLATE, MD_OUTPUT, context)

    print("Converting Markdown → PDF...")
    convert_md_to_pdf(MD_OUTPUT, PDF_OUTPUT)

    print("Done!")
    print(f"Markdown: {os.path.abspath(MD_OUTPUT)}")
    print(f"PDF:      {os.path.abspath(PDF_OUTPUT)}")




def generate_report(title, sections, format="markdown"):

    report = f"# {title}\n\n" if format == "markdown" else f"{title}\n\n"

    for sec in sections:
        heading = sec["heading"]
        content = sec["content"]

        if format == "markdown":
            report += f"## {heading}\n\n{content}\n\n"
        else:
            report += f"{heading}\n{content}\n\n"

    return {
        "report_text": report,
        "created_at": datetime.now().isoformat()
    }


def build_report(
    title: str,
    query: str,
    summary: str,
    laws: list,
    cases: list,
    analysis: str,
    final_answer: str
):
    return {
        "title": title,
        "query": query,
        "summary": summary,
        "laws": laws,
        "cases": cases,
        "analysis": analysis,
        "final_answer": final_answer,
    }

def generate_markdown(report: dict) -> str:
    md = []
    md.append(f"# {report['title']}\n")

    md.append("## 1. 질의 사항")
    md.append(report['query'] + "\n")

    md.append("## 2. 관련 법령")
    for law in report['laws']:
        md.append(f"- {law}")
    md.append("")

    md.append("## 3. 관련 사례")
    for case in report['cases']:
        md.append(f"- {case}")
    md.append("")

    md.append("## 4. 답변 사항")
    md.append(report['analysis'] + "\n")

    md.append("## 최종 결론")
    md.append(report['final_answer'] + "\n")

    return "\n".join(md)

def save_markdown(md_text, filename="report.md"):
    with open(filename, "w", encoding="utf-8") as f:
        f.write(md_text)

from docx import Document

def generate_docx(report: dict, filename="report.docx"):
    doc = Document()

    doc.add_heading(report["title"], level=1)

    doc.add_heading("1. 질의 사항", level=2)
    doc.add_paragraph(report["query"])

    doc.add_heading("2. 관련 법령", level=2)
    for law in report["laws"]:
        doc.add_paragraph(f"- {law}")

    doc.add_heading("3. 관련 사례", level=2)
    for case in report["cases"]:
        doc.add_paragraph(f"- {case}")

    doc.add_heading("4. 답변 사항", level=2)
    doc.add_paragraph(report["analysis"])

    doc.add_heading("최종 결론", level=2)
    doc.add_paragraph(report["final_answer"])

    doc.save(filename)

def generate_docx_from_template(report: dict, output="report.docx"):
    doc = DocxTemplate("report-template.docx")
    doc.render(report)
    doc.save(output)

def generate_pdf(report: dict, filename="report.pdf"):
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(f"<b>{report['title']}</b>", styles["Title"]))
    story.append(Spacer(1, 20))

    def add_section(title, content):
        story.append(Paragraph(f"<b>{title}</b>", styles["Heading2"]))
        if isinstance(content, list):
            for item in content:
                story.append(Paragraph(f"- {item}", styles["Normal"]))
        else:
            story.append(Paragraph(content, styles["Normal"]))
        story.append(Spacer(1, 15))

    add_section("1. 질의 사항", report["query"])
    add_section("2. 관련 법령", report["laws"])
    add_section("3. 관련 사례", report["cases"])
    add_section("4. 답변 사항", report["analysis"])
    add_section("최종 결론", report["final_answer"])

    doc.build(story)


context = {
    "title": "근로기준법 적용 여부 검토",
    "created_at": "2025-12-04",
    "author": "홍길동",
    "reviewer": "김검토",
    "query": "동거 친족에게 근로기준법이 적용되는가?",
    "laws": ["근로기준법 제11조", "고용노동부 행정해석 ○○○"],
    "cases": ["사례번호 2020-15", "국민신문고 유사사례 2022-01"],
    "analysis": "분석 내용...",
    "final_answer": "최종 결론...",
}


doc = DocxTemplate("company_template.docx")
doc.render(context)
doc.save("final_report.docx")

import win32com.client as win32

word = win32.Dispatch('Word.Application')
doc = word.Documents.Open(r"C:\path\input.docx")
doc.SaveAs(r"C:\path\output.pdf", FileFormat=17)  # 17 = PDF
doc.Close()
word.Quit()



